#!/usr/bin/env python3
"""
generate_report_docx.py — One-time DocX report from EDA outputs.

Usage:
    export DATA_PIPELINE_DATA_ROOT=/path/to/runtime
    python scripts/generate_report_docx.py
"""

from __future__ import annotations

import os
import sys
from pathlib import Path

from docx import Document
from docx.shared import Inches, Pt
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH

DATA_ROOT = Path(os.environ.get("DATA_PIPELINE_DATA_ROOT", ""))
if not DATA_ROOT.exists():
    print("ERROR: DATA_PIPELINE_DATA_ROOT is not set or does not exist")
    sys.exit(1)

OUTPUT_DIR = DATA_ROOT / "data-pipeline" / "eda" / "eda-assignment-2" / "output"
REPORT_PATH = OUTPUT_DIR / "eda_report.docx"

GROWER_DATA = [
    ("IL", "DeKalb", "10", "226-1,385", "962", "35%", "846 mm", "1,800"),
    ("IA", "Story", "10", "299-1,152", "735", "31%", "813 mm", "2,139"),
    ("NE", "Phelps", "10", "587-1,182", "787", "25%", "565 mm", "2,241"),
]


def add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    return h


def add_para(doc, text, bold=False, italic=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(11)
    return p


def add_image_section(doc, filename, caption):
    path = OUTPUT_DIR / filename
    if path.exists():
        add_para(doc, caption, italic=True)
        doc.add_picture(str(path), width=Inches(5.5))
        doc.add_paragraph()


def main():
    doc = Document()

    # Title
    title = doc.add_heading("Assignment 2 — Field-Level Exploratory Data Analysis", level=0)
    doc.add_paragraph(
        "Three Corn Belt growers | Real OSM boundaries | NASA POWER weather 2021-2025 | USDA CDL 2021-2025"
    ).runs[0].font.size = Pt(12)
    doc.add_paragraph("Soil analysis excluded per assignment scope.").italic = True

    # Dataset scope
    add_heading(doc, "1. Dataset Scope")
    add_para(doc, "Three growers with 10 fields each, all fields >200 acres for structural comparability.")
    table = doc.add_table(rows=4, cols=8)
    table.style = "Light Shading Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    headers = ["Grower", "County", "Fields", "Range (ac)", "Mean (ac)", "CV", "Avg Precip", "Avg GDD"]
    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = h
    for r, row_data in enumerate(GROWER_DATA):
        for c, val in enumerate(row_data):
            table.rows[r + 1].cells[c].text = val

    # Data sources
    add_heading(doc, "2. Data Sources")
    add_para(doc, "Field boundaries: OpenStreetMap / Overpass API (real OSM, >200ac filter)")
    add_para(doc, "Weather: NASA POWER Zarr grid, ~0.5 degree resolution, daily 2021-2025")
    add_para(doc, "CDL cropland: USDA NASS Cropland Data Layer, 30m resolution, annual 2021-2025")

    # Analysis levels
    add_heading(doc, "3. Analysis Levels")
    add_para(doc, "Within a field: Weather variables (precipitation, GDD) across 5 years for individual fields.")
    add_para(doc, "Across fields within a grower: Field size distributions and corn year frequency within each farm.")
    add_para(doc, "Across growers: IL vs IA vs NE climate, scale, and crop composition comparisons.")

    # Field Boundaries
    add_heading(doc, "4. Field Boundaries")
    add_para(doc, "Statistical viz 1 — Area Histogram: Distribution of field areas across all 30 fields. IL and NE range 200-1,400 ac; IA is slightly tighter. All structurally comparable with CV 25-35%.", bold=False)
    add_image_section(doc, "area_histogram.png", "Fig 1: Area histogram by grower")
    add_para(doc, "Statistical viz 2 — Area Boxplot: Side-by-side comparison of field area medians (~960 IL, ~720 IA, ~780 NE). Overlapping ranges confirm fair comparison.", bold=False)
    add_image_section(doc, "area_boxplot.png", "Fig 2: Area boxplot by grower")
    add_para(doc, "Comparison — Area comparison table (CSV): n, min, max, mean, CV, total per grower.", bold=False)
    add_para(doc, "Map — Field boundaries colored by grower: Three counties in three states. Spatial context.", bold=False)
    add_image_section(doc, "boundary_map.png", "Fig 3: Field boundaries by grower")

    # Weather
    add_heading(doc, "5. Weather")
    add_para(doc, "Statistical viz 1 — Annual Precipitation: Grouped bars by grower and year. NE receives ~300mm less rain than IL. 2022 drought visible in NE (392mm vs 565mm avg).", bold=False)
    add_image_section(doc, "annual_precip.png", "Fig 4: Annual precipitation by grower")
    add_para(doc, "Statistical viz 2 — Annual GDD: NE has ~400 more growing degree days than IL. Inverse of precipitation: wet=cool, dry=hot.", bold=False)
    add_image_section(doc, "annual_gdd.png", "Fig 5: Annual GDD by grower")
    add_para(doc, "Comparison — Precip vs GDD Scatter: 15 points (3 growers x 5 years). IL clusters wet/cool, NE clusters dry/hot. The fundamental climate tradeoff.", bold=False)
    add_image_section(doc, "precip_gdd_correlation.png", "Fig 6: Precip vs GDD correlation")
    add_para(doc, "Map — Weather centroids: Field centroids colored by grower. Confirms three distinct climate zones.", bold=False)
    add_image_section(doc, "weather_centroid_map.png", "Fig 7: Weather centroid map")

    # CDL
    add_heading(doc, "6. CDL / Cropland Data")
    add_para(doc, "Statistical viz 1 — Crop Composition Stacked Bar: NE is corn-dominant, IA leans soy, IL is balanced with more alfalfa and wheat.", bold=False)
    add_image_section(doc, "crop_composition_stacked.png", "Fig 8: Crop composition by grower")
    add_para(doc, "Statistical viz 2 — Corn Years Histogram: NE has 2 fields with 5 years continuous corn (irrigation). IL and IA max at 4. Counts dominant crop per field-year.", bold=False)
    add_image_section(doc, "corn_years_histogram.png", "Fig 9: Corn years histogram")
    add_para(doc, "Comparison — Corn/Soy Ratio: Line chart of combined corn+soy % over time. All three >75% every year. NE corn-dominant, IA balanced, IL variable.", bold=False)
    add_image_section(doc, "corn_soy_ratio.png", "Fig 10: Corn/soy ratio by grower")
    add_para(doc, "Map — Dominant 2025 Crop: Fields colored by dominant crop. Corn=yellow, soy=green. NE is heavily yellow (corn-dominant).", bold=False)
    add_image_section(doc, "dominant_crop_map.png", "Fig 11: Dominant crop map")

    # Key findings
    add_heading(doc, "7. Key Findings")
    findings = [
        "Fields are structurally comparable. All >200ac, CV 25-35% across states.",
        "Climate gradient drives strategy. NE gets ~300mm less rain than IL but has ~400 more GDD. Wet=cool, dry=hot.",
        "Irrigation enables continuous corn. NE has 2 fields with 5 years of corn (out of 5) because irrigation buffers drought risk.",
        "2022 drought signal. Nebraska's 2022 precipitation (392mm) was 31% below its 5-year mean. No CDL yield penalty — consistent with irrigation.",
        "Without yield data, this EDA describes planting strategy, not productivity. We can show what farmers plant, not how well it performs.",
    ]
    for f in findings:
        doc.add_paragraph(f, style="List Bullet")

    # Limitations
    add_heading(doc, "8. Limitations")
    limitations = [
        "Small sample: 10 fields per state from one county each.",
        "Short time window: 5 years insufficient for climate trend conclusions.",
        "Modeled weather: NASA POWER reanalysis (~50km grid), not on-site stations.",
        "Provisional 2025 CDL: Preliminary, may be revised by USDA.",
        "OSM boundary quality: Crowd-sourced, may have omissions or grouping issues.",
        "Single county per state: Cannot capture within-state diversity.",
        "Soil excluded per assignment scope.",
        "No yield data: EDA describes planting strategy, not productivity outcomes.",
    ]
    for lim in limitations:
        doc.add_paragraph(lim, style="List Bullet")

    # Save
    REPORT_PATH.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(REPORT_PATH))
    print(f"Report saved: {REPORT_PATH}")


if __name__ == "__main__":
    main()
